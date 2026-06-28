"""File format parsers for DOCX, JSON, and PDF documents."""

import io
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional
import pdfplumber
from docx import Document as DocxDocument
from ingestion.scrapers.base_scraper import RawDocument

logger = logging.getLogger("ingestion.processors.parsers")


class FileParser:
    """Base class for file parsers."""

    async def parse(self, file_content: bytes, filename: str, institution_slug: str) -> RawDocument:
        """Parse file content and return a RawDocument.
        
        Args:
            file_content: Binary file content
            filename: Name of the file for reference
            institution_slug: Slug of the institution this document belongs to
            
        Returns:
            RawDocument with parsed content
        """
        raise NotImplementedError


class DocxParser(FileParser):
    """Parser for Microsoft Word (.docx) files."""

    async def parse(self, file_content: bytes, filename: str, institution_slug: str) -> RawDocument:
        """Extract text and tables from DOCX file."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(row_cells))
                if table_rows:
                    text_parts.append("\n".join(table_rows))
            
            full_text = "\n\n".join(text_parts)
            
            if not full_text.strip():
                logger.warning(f"DOCX file {filename} produced empty content")
                full_text = "(Empty document)"
            
            return RawDocument(
                url=f"file://{filename}",
                category="uploaded_document",
                sub_category="docx",
                institution_slug=institution_slug,
                title=Path(filename).stem,
                http_status=200,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                raw_text=full_text,
                scraped_at=datetime.utcnow(),
            )
        except Exception as e:
            logger.error(f"Failed to parse DOCX file {filename}: {e}", exc_info=True)
            raise


class JSONParser(FileParser):
    """Parser for JSON files."""

    async def parse(self, file_content: bytes, filename: str, institution_slug: str) -> RawDocument:
        """Parse JSON file and convert to structured text."""
        try:
            data = json.loads(file_content.decode("utf-8"))
            text_lines = self._flatten_json(data, key_prefix="")
            full_text = "\n".join(text_lines)
            
            if not full_text.strip():
                logger.warning(f"JSON file {filename} produced empty content")
                full_text = "(Empty JSON)"
            
            return RawDocument(
                url=f"file://{filename}",
                category="uploaded_document",
                sub_category="json",
                institution_slug=institution_slug,
                title=Path(filename).stem,
                http_status=200,
                content_type="application/json",
                raw_text=full_text,
                scraped_at=datetime.utcnow(),
            )
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON file {filename}: Invalid JSON - {e}", exc_info=True)
            raise ValueError(f"Invalid JSON in file {filename}") from e
        except Exception as e:
            logger.error(f"Failed to parse JSON file {filename}: {e}", exc_info=True)
            raise

    def _flatten_json(self, obj: Any, key_prefix: str = "") -> List[str]:
        """Recursively flatten JSON object into key-value text lines."""
        lines = []
        
        if isinstance(obj, dict):
            for key, value in obj.items():
                new_prefix = f"{key_prefix}.{key}" if key_prefix else key
                if isinstance(value, (dict, list)):
                    lines.extend(self._flatten_json(value, new_prefix))
                else:
                    # Format key-value pair
                    if value is not None:
                        lines.append(f"{new_prefix}: {str(value)}")
        elif isinstance(obj, list):
            for idx, item in enumerate(obj):
                new_prefix = f"{key_prefix}[{idx}]"
                if isinstance(item, (dict, list)):
                    lines.extend(self._flatten_json(item, new_prefix))
                else:
                    if item is not None:
                        lines.append(f"{new_prefix}: {str(item)}")
        else:
            if obj is not None:
                lines.append(f"{key_prefix}: {str(obj)}")
        
        return lines


class PDFParser(FileParser):
    """Parser for PDF files."""

    async def parse(self, file_content: bytes, filename: str, institution_slug: str) -> RawDocument:
        """Extract text and tables from PDF file."""
        try:
            text_parts = []
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                    
                    # Extract tables
                    try:
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                table_text = self._table_to_text(table)
                                if table_text:
                                    text_parts.append(f"--- Table (Page {page_num}) ---\n{table_text}")
                    except Exception as e:
                        logger.debug(f"Failed to extract tables from page {page_num}: {e}")
            
            full_text = "\n\n".join(text_parts)
            
            if not full_text.strip():
                logger.warning(f"PDF file {filename} produced empty content")
                full_text = "(Empty PDF)"
            
            return RawDocument(
                url=f"file://{filename}",
                category="uploaded_document",
                sub_category="pdf",
                institution_slug=institution_slug,
                title=Path(filename).stem,
                http_status=200,
                content_type="application/pdf",
                raw_text=full_text,
                scraped_at=datetime.utcnow(),
            )
        except Exception as e:
            logger.error(f"Failed to parse PDF file {filename}: {e}", exc_info=True)
            raise

    def _table_to_text(self, table: List[List[Optional[str]]]) -> str:
        """Convert table grid to text representation."""
        if not table or not table[0]:
            return ""
        
        lines = []
        for row in table:
            row_text = " | ".join(str(cell or "").strip() for cell in row)
            lines.append(row_text)
        
        return "\n".join(lines)


class UniversalFileParser(FileParser):
    """Routes file parsing to appropriate parser based on file extension."""

    def __init__(self) -> None:
        self.parsers: Dict[str, FileParser] = {
            ".docx": DocxParser(),
            ".doc": DocxParser(),
            ".json": JSONParser(),
            ".pdf": PDFParser(),
        }

    async def parse(self, file_content: bytes, filename: str, institution_slug: str) -> RawDocument:
        """Parse file using the appropriate parser."""
        file_ext = Path(filename).suffix.lower()
        
        if file_ext not in self.parsers:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                f"Supported formats: {', '.join(self.parsers.keys())}"
            )
        
        parser = self.parsers[file_ext]
        logger.info(f"Parsing {file_ext} file: {filename} for institution: {institution_slug}")
        return await parser.parse(file_content, filename, institution_slug)
